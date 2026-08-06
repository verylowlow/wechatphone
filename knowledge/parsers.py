"""File parsers: raw file -> plain text.

Supported: txt / md / markdown / html / htm / csv / pdf / docx.
"""
from __future__ import annotations

import csv
import io
import re

SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".html", ".htm", ".csv", ".pdf", ".docx",
}


class ParseError(Exception):
    pass


def parse_file(path: str) -> str:
    """Parse a file into plain text. Raises ParseError on failure."""
    ext = _ext(path)
    if ext not in SUPPORTED_EXTENSIONS:
        raise ParseError(f"unsupported file type: {ext}")
    try:
        if ext in (".txt", ".md", ".markdown"):
            return _read_text(path)
        if ext in (".html", ".htm"):
            return _parse_html(path)
        if ext == ".csv":
            return _parse_csv(path)
        if ext == ".pdf":
            return _parse_pdf(path)
        if ext == ".docx":
            return _parse_docx(path)
    except ParseError:
        raise
    except Exception as e:  # noqa: BLE001
        raise ParseError(f"failed to parse {path}: {e}") from e
    raise ParseError(f"unsupported file type: {ext}")


def _ext(path: str) -> str:
    import os
    return os.path.splitext(path)[1].lower()


def _read_text(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ParseError(f"cannot decode text file: {path}")


def _parse_html(path: str) -> str:
    with open(path, "rb") as f:
        raw = f.read().decode("utf-8", errors="ignore")
    raw = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", raw)
    text = re.sub(r"(?s)<[^>]+>", "\n", raw)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&lt;", "<", text)
    text = re.sub(r"&gt;", ">", text)
    text = re.sub(r"&amp;", "&", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _parse_csv(path: str) -> str:
    text = _read_text(path)
    rows = list(csv.reader(io.StringIO(text)))
    if not rows:
        return ""
    header = rows[0]
    lines = []
    for row in rows[1:]:
        pairs = [f"{h}: {v}" for h, v in zip(header, row) if v.strip()]
        if pairs:
            lines.append("；".join(pairs))
    return "\n".join(lines)


def _parse_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages):
        t = page.extract_text() or ""
        if t.strip():
            pages.append(t)
    if not pages:
        raise ParseError("PDF contains no extractable text (scanned image?)")
    return "\n\n".join(pages)


def _parse_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("；".join(cells))
    return "\n".join(parts)
